"""RAG knowledge management API: upload, list, delete documents."""

import logging
import os
import uuid
import datetime
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session
from app.core.response import UTF8JSONResponse
from app.database import get_db
from app.models.models import KnowledgeDocument
from app.services.rag_service import ingest_documents, query_knowledge
from app.services.code_rag_service import ingest_code_file, is_supported_code_file, query_codebase

router = APIRouter(prefix="/api/rag", tags=["rag"], default_response_class=UTF8JSONResponse)

UPLOAD_DIR = Path(__file__).resolve().parents[2] / "data" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx", ".py", ".js", ".jsx", ".java"}
SUPPORTED_CODE_LANGUAGES = {"python", "javascript", "java"}
logger = logging.getLogger(__name__)



class CodeSearchRequest(BaseModel):
    """JSON body for codebase semantic search."""

    query: str
    top_k: int = Field(default=5, ge=1, le=20)
    language: str | None = None
    file_name: str | None = None

def _extract_text(file_path: Path, suffix: str) -> str:
    """Extract plain text from supported file formats."""
    if suffix in (".md", ".txt", ".py", ".js", ".jsx", ".java"):
        return file_path.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(file_path))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF parse failed: {e}")

    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(file_path))
            return "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOCX parse failed: {e}")

    raise HTTPException(status_code=400, detail=f"Unsupported file extension: {suffix}")


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    category: str = Form("custom"),
    tags: str = Form(""),
    db: Session = Depends(get_db),
):
    """Upload a document or supported code file into the matching RAG index."""
    suffix = Path(file.filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported extension {suffix}; allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    # Save uploaded file
    safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
    save_path = UPLOAD_DIR / safe_name
    content = await file.read()
    save_path.write_bytes(content)

    # Extract text
    text = _extract_text(save_path, suffix)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Uploaded file is empty or could not be parsed")

    # Derive title
    title = Path(file.filename).stem
    is_code = is_supported_code_file(file.filename)

    # Save to DB
    doc = KnowledgeDocument(
        title=title,
        source=f"uploads/{safe_name}",
        category=category,
        doc_type="code" if is_code else "user_upload",
        content=text,
        tags=tags or title,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # Ingest into the matching ChromaDB collection.
    if is_code:
        ingest_result = ingest_code_file(save_path, original_name=file.filename)
    else:
        ingest_result = ingest_documents()

    return {
        "code": 0,
        "data": {
            "id": doc.id,
            "title": doc.title,
            "filename": file.filename,
            "size": len(content),
            "chars": len(text),
            "category": category,
            "doc_type": doc.doc_type,
            "ingest": ingest_result,
        },
        "message": "ok",
    }


@router.get("/documents")
def list_documents(db: Session = Depends(get_db)):
    """List all knowledge documents."""
    docs = db.query(KnowledgeDocument).order_by(KnowledgeDocument.created_at.desc()).all()
    return {
        "code": 0,
        "data": [
            {
                "id": d.id,
                "title": d.title,
                "source": d.source,
                "category": d.category,
                "doc_type": d.doc_type,
                "tags": d.tags,
                "chars": len(d.content or ""),
                "created_at": str(d.created_at),
            }
            for d in docs
        ],
        "message": "ok",
    }


@router.delete("/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    """Delete a knowledge document."""
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Delete physical file if it is an upload.
    if doc.source and doc.source.startswith("uploads/"):
        file_path = Path(__file__).resolve().parents[2] / "data" / doc.source
        if file_path.exists():
            file_path.unlink()

    db.delete(doc)
    db.commit()

    # Keep the legacy SOP index behavior unchanged.
    ingest_documents()

    return {"code": 0, "data": {"id": doc_id}, "message": "Document deleted"}


def _build_code_search_where(language: str | None, file_name: str | None) -> dict | None:
    conditions = []
    if language:
        normalized_language = language.strip().lower()
        if normalized_language not in SUPPORTED_CODE_LANGUAGES:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported language: {language}. Allowed: {', '.join(sorted(SUPPORTED_CODE_LANGUAGES))}",
            )
        conditions.append({"language": normalized_language})

    if file_name:
        normalized_file_name = file_name.strip()
        if normalized_file_name:
            conditions.append({"file_name": normalized_file_name})

    if not conditions:
        return None
    if len(conditions) == 1:
        return conditions[0]
    return {"$and": conditions}


@router.post("/code-search")
def search_codebase(payload: CodeSearchRequest):
    """Search the indexed codebase collection and return code chunks with metadata."""
    query = payload.query.strip()
    if not query:
        raise HTTPException(status_code=400, detail="query is required")

    where = _build_code_search_where(payload.language, payload.file_name)
    try:
        results = query_codebase(query=query, top_k=payload.top_k, where=where)
    except (ImportError, ModuleNotFoundError) as exc:
        logger.exception("Code RAG dependency unavailable for query=%r", query)
        raise HTTPException(status_code=503, detail="Code RAG backend is unavailable") from exc
    except Exception as exc:
        module_name = exc.__class__.__module__
        if module_name.startswith("chromadb"):
            logger.exception("ChromaDB unavailable for code query=%r", query)
            raise HTTPException(status_code=503, detail="Code RAG backend is unavailable") from exc
        logger.exception("Code RAG search failed for query=%r", query)
        raise HTTPException(status_code=500, detail="Code RAG search failed") from exc

    results = sorted(results, key=lambda item: item.get("score", 0), reverse=True)
    logger.info("Code RAG search query=%r hits=%d", query, len(results))
    return {
        "code": 0,
        "message": "ok",
        "data": {
            "query": query,
            "total": len(results),
            "results": results,
        },
    }

@router.post("/search")
def search_knowledge(
    query: str = Form(...),
    top_k: int = Form(4),
):
    """Search the RAG knowledge base."""
    results = query_knowledge(query, top_k=top_k)
    return {"code": 0, "data": results, "message": "ok"}
