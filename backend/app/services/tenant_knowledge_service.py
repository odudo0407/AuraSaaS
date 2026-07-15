"""Tenant-scoped private knowledge RAG service."""

from __future__ import annotations

import hashlib
import logging
import re
import uuid
from pathlib import Path

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.models import TenantKnowledgeDocument
from app.services.rag_service import _chunk_text

logger = logging.getLogger(__name__)

COLLECTION_NAME = "aurasaas_tenant_knowledge"
ROOT_DIR = Path(__file__).resolve().parents[3]
UPLOAD_ROOT = ROOT_DIR / "backend" / "data" / "uploads" / "knowledge"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def _tenant_id_value(tenant_id: int | str) -> str:
    return str(tenant_id).strip()


def _where_for_file(tenant_id: int | str, file_name: str) -> dict:
    return {"$and": [{"tenant_id": _tenant_id_value(tenant_id)}, {"source_file": file_name}]}


def _where_for_tenant(tenant_id: int | str) -> dict:
    return {"tenant_id": _tenant_id_value(tenant_id)}


def _embedding_function():
    from chromadb.utils import embedding_functions

    settings = get_settings()
    cache_dir = Path(settings.code_embedding_cache_dir)
    cache_dir.mkdir(parents=True, exist_ok=True)
    return embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name=settings.code_embedding_model,
        cache_folder=str(cache_dir),
    )


def _get_collection():
    import chromadb

    settings = get_settings()
    client = chromadb.PersistentClient(path=settings.chroma_dir)
    return client.get_or_create_collection(
        COLLECTION_NAME,
        embedding_function=_embedding_function(),
        metadata={"hnsw:space": "cosine"},
    )


def _delete_vectors_for_file(tenant_id: int | str, file_name: str) -> None:
    collection = _get_collection()
    try:
        collection.delete(where=_where_for_file(tenant_id, file_name))
    except (ValueError, KeyError) as exc:
        logger.debug("No tenant vectors to delete tenant_id=%s file=%s: %s", tenant_id, file_name, exc)
    except Exception as exc:
        if exc.__class__.__module__.startswith("chromadb"):
            logger.debug("Ignoring Chroma delete miss tenant_id=%s file=%s: %s", tenant_id, file_name, exc)
            return
        raise


def _save_upload(file_bytes: bytes, tenant_id: int | str, file_name: str) -> Path:
    tenant_dir = UPLOAD_ROOT / _tenant_id_value(tenant_id)
    tenant_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"{uuid.uuid4().hex[:8]}_{Path(file_name).name}"
    file_path = tenant_dir / safe_name
    file_path.write_bytes(file_bytes)
    return file_path


def _extract_pdf(file_path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(file_path: Path) -> str:
    from docx import Document

    doc = Document(str(file_path))
    return "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())


def extract_text(file_path: Path, file_type: str) -> str:
    """Extract raw text from a supported tenant knowledge file."""

    suffix = file_type.lower()
    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding=get_settings().text_encoding, errors="ignore")
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)
    if suffix == ".xlsx":
        raise ValueError("Excel knowledge upload is planned for Phase 2")
    raise ValueError(f"Unsupported tenant knowledge file type: {suffix}")


def clean_text(text: str) -> str:
    """Normalize tenant document text for chunking and retrieval."""

    text = text.replace("\x00", "")
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _build_chunk_ids(doc_id: int, chunks: list[str]) -> list[str]:
    return [f"tenant-doc-{doc_id}-{idx}" for idx, _ in enumerate(chunks)]


def _add_chunks(doc: TenantKnowledgeDocument, chunks: list[str]) -> None:
    collection = _get_collection()
    tenant_id = _tenant_id_value(doc.tenant_id)
    ids = _build_chunk_ids(doc.id, chunks)
    metadatas = [
        {
            "scope": "tenant",
            "tenant_id": tenant_id,
            "doc_id": doc.id,
            "source_file": doc.file_name,
            "chunk_index": idx,
            "file_type": doc.file_type,
        }
        for idx, _ in enumerate(chunks)
    ]
    if ids:
        collection.add(ids=ids, documents=chunks, metadatas=metadatas)


def _delete_old_documents(db: Session, tenant_id: int, file_name: str) -> None:
    old_docs = db.query(TenantKnowledgeDocument).filter(
        TenantKnowledgeDocument.tenant_id == tenant_id,
        TenantKnowledgeDocument.file_name == file_name,
    ).all()
    for old_doc in old_docs:
        _delete_vectors_for_file(tenant_id, old_doc.file_name)
        if old_doc.source_path:
            try:
                Path(old_doc.source_path).unlink(missing_ok=True)
            except OSError:
                logger.warning("Failed to delete old tenant knowledge file: %s", old_doc.source_path)
        db.delete(old_doc)
    if old_docs:
        db.flush()


def ingest_tenant_knowledge_file(
    db: Session,
    *,
    tenant_id: int,
    file_name: str,
    file_bytes: bytes,
    created_by: str | None = None,
) -> TenantKnowledgeDocument:
    """Save, parse, chunk, and index a tenant private knowledge document."""

    suffix = Path(file_name).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported tenant knowledge file type: {suffix}")
    if not file_bytes:
        raise ValueError("Uploaded file is empty")

    content_hash = hashlib.sha256(file_bytes).hexdigest()
    _delete_old_documents(db, tenant_id, file_name)
    file_path = _save_upload(file_bytes, tenant_id, file_name)

    doc = TenantKnowledgeDocument(
        tenant_id=tenant_id,
        file_name=Path(file_name).name,
        file_size=len(file_bytes),
        file_type=suffix,
        source_path=str(file_path),
        content_hash=content_hash,
        chunk_count=0,
        status="processing",
        created_by=created_by or "",
    )
    db.add(doc)
    db.flush()

    try:
        raw_text = extract_text(file_path, suffix)
        cleaned = clean_text(raw_text)
        if not cleaned:
            raise ValueError("No text could be extracted from uploaded file")

        chunks = [chunk for chunk in _chunk_text(cleaned) if chunk.strip()]
        _delete_vectors_for_file(tenant_id, doc.file_name)
        _add_chunks(doc, chunks)
        doc.chunk_count = len(chunks)
        doc.status = "completed"
        doc.error_message = None
        db.commit()
        db.refresh(doc)
        logger.info("Tenant knowledge ingested tenant_id=%s file=%s chunks=%d", tenant_id, doc.file_name, len(chunks))
        return doc
    except Exception as exc:
        doc.status = "failed"
        doc.error_message = str(exc)[:2000]
        db.commit()
        db.refresh(doc)
        logger.exception("Tenant knowledge ingest failed tenant_id=%s file=%s", tenant_id, file_name)
        raise


def query_tenant_knowledge(query: str, tenant_id: int | str, top_k: int = 4) -> list[dict]:
    """Query private tenant knowledge with mandatory tenant isolation."""

    collection = _get_collection()
    result = collection.query(query_texts=[query], n_results=top_k, where=_where_for_tenant(tenant_id))
    docs = result.get("documents", [[]])[0]
    metas = result.get("metadatas", [[]])[0]
    distances = result.get("distances", [[]])[0]
    return [
        {
            "title": meta.get("source_file", "Tenant knowledge"),
            "snippet": doc[:500],
            "score": float(1 / (1 + distance)) if distance is not None else 0.0,
            "source": meta.get("source_file", ""),
            "category": "tenant_knowledge",
            "tags": "tenant,private",
            "type": "tenant_knowledge",
            "metadata": meta,
        }
        for doc, meta, distance in zip(docs, metas, distances)
    ]


def delete_tenant_knowledge_document(db: Session, *, doc_id: int, tenant_id: int) -> TenantKnowledgeDocument | None:
    """Delete a tenant document record, vectors, and best-effort physical file."""

    doc = db.query(TenantKnowledgeDocument).filter(
        TenantKnowledgeDocument.id == doc_id,
        TenantKnowledgeDocument.tenant_id == tenant_id,
    ).first()
    if doc is None:
        return None

    _delete_vectors_for_file(tenant_id, doc.file_name)
    source_path = doc.source_path
    db.delete(doc)
    db.commit()
    if source_path:
        try:
            Path(source_path).unlink(missing_ok=True)
        except OSError:
            logger.warning("Failed to delete tenant knowledge file: %s", source_path)
    return doc